from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "practice" / "בחינה לדוגמה - הנדסת נתונים - טיוטה 4.docx"

HELPERS_PATH = ROOT / "scripts" / "build_practice4_docs.py"
spec = importlib.util.spec_from_file_location("build_practice4_docs", HELPERS_PATH)
if spec is None or spec.loader is None:
    raise RuntimeError(f"Could not load DOCX helpers from {HELPERS_PATH}")
helpers = importlib.util.module_from_spec(spec)
spec.loader.exec_module(helpers)

configure_document = helpers.configure_document
create_list_numbering = helpers.create_list_numbering
set_cell_margins = helpers.set_cell_margins
set_cell_width = helpers.set_cell_width
set_paragraph_direction = helpers.set_paragraph_direction
set_paragraph_numbering = helpers.set_paragraph_numbering
set_run_direction = helpers.set_run_direction
set_run_size = helpers.set_run_size
set_table_direction = helpers.set_table_direction
set_table_width = helpers.set_table_width
shade_cell = helpers.shade_cell
cm_to_dxa = helpers.cm_to_dxa


DATASETS = [
    {
        "title": "אוסף נתונים א׳ - כרטיסי תמיכה",
        "description": (
            "נתון אוסף כרטיסי תמיכה קצר. מזהה הכרטיס כבר הוסר, והטקסט מוצג כרשימת טוקנים. "
            "אוסף זה ישמש בשאלות TF-IDF וחיפוש סמנטי."
        ),
        "tables": [
            {
                "headers": ["doc_id", "tokens", "אורך מסמך"],
                "rows": [
                    ["D1", "vpn, not, working, today", "4"],
                    ["D2", "vpn, not, working, now", "4"],
                    ["D3", "vpn, not, responding, today", "4"],
                    ["D4", "printer, working, fine, now", "4"],
                    ["D5", "printer, not, working", "3"],
                ],
            },
            {
                "headers": ["item", "vector"],
                "rows": [
                    ["Q", "[0.85, 0.65, 0.05]"],
                    ["C1", "[0.85, 0.45, 0.05]"],
                    ["C2", "[0.80, 0.75, 0.02]"],
                    ["C3", "[0.60, 0.90, 0.03]"],
                    ["C4", "[0.20, 0.85, 0.25]"],
                    ["C5", "[0.05, 0.35, 0.90]"],
                ],
            },
        ],
    },
    {
        "title": "אוסף נתונים ב׳ - קטעי מדיניות למערכת RAG",
        "description": (
            "נתון מנגנון Retrieval עבור עוזר קורס. השאילתה היא: late homework penalty. "
            "עבור BM25 השתמשו בפרמטרים avgdl=6, k1=1.2, b=0.75 ובדף הנוסחאות."
        ),
        "tables": [
            {
                "headers": ["chunk", "matched query terms", "length"],
                "rows": [
                    ["C1", "late, homework, penalty", "8"],
                    ["C2", "none", "5"],
                    ["C3", "homework", "6"],
                ],
            },
            {
                "headers": ["term", "IDF"],
                "rows": [
                    ["late", "0.981"],
                    ["homework", "0.470"],
                    ["penalty", "0.981"],
                ],
            },
            {
                "headers": ["chunk", "BM25 normalized", "cosine", "tokens"],
                "rows": [
                    ["C1", "0.780", "0.986", "40"],
                    ["C2", "0.950", "0.995", "40"],
                    ["C3", "0.820", "0.946", "40"],
                    ["C4", "0.550", "0.755", "40"],
                    ["C5", "0.180", "0.304", "30"],
                ],
            },
        ],
    },
    {
        "title": "אוסף נתונים ג׳ - אירועי צפייה וגרף קישורים",
        "description": (
            "נתונים אלה ישמשו בשאלת חישוב מבוזר ו-PageRank. עבור PageRank: העמודים הם A, B, C, D; "
            "הדירוג ההתחלתי של כל עמוד הוא 0.25; d=0.85; והעמוד D הוא dangling node."
        ),
        "tables": [
            {
                "headers": ["event_id", "category", "product_id"],
                "rows": [
                    ["1", "books", "B1"],
                    ["2", "books", "B1"],
                    ["3", "books", "B2"],
                    ["4", "games", "G1"],
                    ["5", "games", "G2"],
                    ["6", "games", "G2"],
                ],
            },
            {
                "headers": ["from_page", "to_page"],
                "rows": [
                    ["A", "B"],
                    ["A", "C"],
                    ["B", "C"],
                    ["C", "A"],
                ],
            },
        ],
    },
    {
        "title": "אוסף נתונים ד׳ - מקורות למחסן נתונים",
        "description": (
            "נתונים תפעוליים מגיעים ממספר מערכות מקור. יש לתכנן מחסן נתונים ותהליך טעינה שמתאימים "
            "לדוחות מכירות, החזרות וניתוח לקוחות."
        ),
        "tables": [
            {
                "headers": ["source", "important columns"],
                "rows": [
                    ["orders", "order_id, customer_id, store_id, order_date, status, updated_at"],
                    ["order_items", "order_id, product_id, quantity, unit_price"],
                    ["products", "product_id, category, brand"],
                    ["customers", "customer_id, full_name, email, birth_date, status"],
                    ["returns", "order_id, product_id, return_date"],
                ],
            },
            {
                "headers": ["source column", "target candidate"],
                "rows": [
                    ["customer_id", "source_customer_id"],
                    ["full_name", "customer_name"],
                    ["birth_date", "age_group"],
                    ["email", "email_domain"],
                    ["status", "is_active"],
                ],
            },
        ],
    },
]


QUESTIONS = [
    {
        "title": "שאלה 1 (25 נקודות)",
        "intro": "השתמשו באוסף נתונים א׳. בשאלה זו יש להציג חישובי עזר ונימוקים.",
        "parts": [
            "(8 נק׳) חשבו את df ואת idf עבור המונחים vpn, today, responding, printer, working. לאחר מכן חשבו TF-IDF עבור vpn ב-D1, today ב-D1, responding ב-D3, printer ב-D5 ו-working ב-D5.",
            "(7 נק׳) השתמשו בווקטור Q ובווקטורים C1, C2, C3. חשבו dot product, נורמה וציון cosine לכל אחד משלושת הקטעים, ודרגו אותם לפי מידת הדמיון לשאילתה.",
            "(5 נק׳) ללא תלות בסעיפים הקודמים, הסבירו מה ישתנה אם במקום unigram TF-IDF נרצה להשתמש ב-bigram TF-IDF עבור הביטוי not working.",
            "(5 נק׳) האם cosine similarity יכול להחזיר קטע רלוונטי גם כאשר אין חפיפה מילולית מלאה בין השאילתה לבין הקטע? הסבירו באמצעות הדוגמה.",
        ],
    },
    {
        "title": "שאלה 2 (25 נקודות)",
        "intro": "השתמשו באוסף נתונים ב׳. השאלה עוסקת ב-BM25, ציון משולב והערכת Retrieval במערכת RAG.",
        "parts": [
            "(8 נק׳) חשבו את ציון BM25 עבור C1, C2, C3. יש להציג את נרמול האורך, תרומת המונחים, ואת הדירוג הסופי.",
            "(6 נק׳) עבור הטבלה עם ציוני BM25 מנורמלים ו-cosine, חשבו ציון משולב שבו BM25 מקבל משקל 40% ו-cosine מקבל משקל 60%. דרגו את חמשת הקטעים.",
            "(5 נק׳) חלון ההקשר הוא 280 טוקנים. רכיבים קבועים: system policy=80, user question=25, answer budget=45, audit reserve=10. האם top-3 נכנס בתקציב הראיות? הציגו חישוב.",
            "(6 נק׳) נתון כי הקטעים הרלוונטיים באמת הם C1, C2, C3. חשבו precision@3 ו-recall@3 עבור הדירוג שקיבלתם, והסבירו האם ניתן לאפשר למודל לייצר תשובה מבוססת מקורות.",
        ],
    },
    {
        "title": "שאלה 3 (25 נקודות) - בשאלה זו הסעיפים בלתי תלויים זה בזה",
        "intro": "השתמשו באוסף נתונים ג׳. יש להציג את שלבי החישוב ולא רק את התוצאה הסופית.",
        "parts": [
            "(7 נק׳) עבור טבלת אירועי הצפייה, כתבו את פלט ה-Mapper, את קבוצות ה-Shuffle, ואת פלט ה-Reducer לספירת צפיות לכל זוג category, product_id.",
            "(5 נק׳) המשיכו את הסעיף הקודם ותארו שלב MapReduce נוסף שמחזיר לכל קטגוריה את המוצר עם מספר הצפיות הגבוה ביותר. האם Combiner מתאים לשלב הספירה? נמקו.",
            "(13 נק׳) עבור גרף הקישורים, בצעו איטרציה אחת של PageRank. חשבו את תרומות הדירוג, dangling mass, סכום התרומות הנכנסות לכל עמוד, ואת הדירוג החדש של A, B, C, D.",
        ],
    },
    {
        "title": "שאלה 4 (25 נקודות)",
        "intro": "השתמשו באוסף נתונים ד׳. השאלה עוסקת בתכנון מחסן נתונים, ETL ו-STTM.",
        "parts": [
            "(8 נק׳) תכננו סכמת star schema לדוחות מכירות. ציינו את טבלת ה-fact, טבלאות ה-dimension, grain, מפתחות, ושני מדדים מרכזיים.",
            "(6 נק׳) כתבו STTM עבור לפחות ארבע עמודות יעד מתוך טבלת הלקוחות. לכל עמודה ציינו מקור, טרנספורמציה, וכלל איכות אחד.",
            "(6 נק׳) תארו תהליך ETL לילי לטעינת fact_sales. התייחסו ל-deduplication, רשומות חסרות, late-arriving records, ו-idempotency.",
            "(5 נק׳) במידה ופרטי הלקוח משתנים לאורך זמן, כיצד הייתם מייצגים זאת במחסן הנתונים? הסבירו מה משתנה ב-STTM ובתהליך הטעינה.",
        ],
    },
]


def add_rtl_paragraph(
    doc: Document,
    text: str = "",
    style: str | None = None,
    *,
    bold: bool = False,
    size: float = 11,
    color: RGBColor | None = None,
    space_after: float | None = None,
):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_direction(paragraph, True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    if text:
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        set_run_size(run, size)
        set_run_direction(run, True)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    add_rtl_paragraph(
        doc,
        text,
        style="Heading 1" if level == 1 else "Heading 2",
        bold=True,
        size=16 if level == 1 else 13,
        color=RGBColor(42, 82, 125),
        space_after=6,
    )


def set_cell_text(cell, text: str, *, bold: bool = False, header: bool = False, width_dxa: int | None = None) -> None:
    if width_dxa is not None:
        set_cell_width(cell, width_dxa)
    set_cell_margins(cell, top=70, bottom=70, start=90, end=90)
    if header:
        shade_cell(cell, "D9EAF7")
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    set_paragraph_direction(paragraph, True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.bold = bold
    set_run_size(run, 9.2 if not header else 9.8)
    set_run_direction(run, True)


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_direction(table, rtl=True)
    table_width = cm_to_dxa(17.2)
    set_table_width(table, table_width)
    column_width = table_width // len(headers)

    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, header=True, width_dxa=column_width)

    for row_values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_values):
            set_cell_text(cell, value, width_dxa=column_width)

    add_rtl_paragraph(doc, space_after=3)


def add_exam_header(doc: Document) -> None:
    lines = [
        "בית הספר להנדסת תעשייה וניהול",
        "מס׳ בחינה_________________",
        "תעודת זהות________________",
        "הנדסת נתונים - בחינה לדוגמה",
        "משך הבחינה: שעתיים וחצי.",
        "חומר עזר: מחשבון ודף נוסחאות שמצורף לבחינה.",
        "מבנה הבחינה:",
        "בבחינה 4 שאלות פתוחות, יש לענות על כולן.",
        "יש לענות על השאלות בטופס הבחינה בלבד.",
        "ניתן להשתמש במחברת בחינה לצורך טיוטה בלבד; מחברת הטיוטה לא תיבדק ולא תיסרק.",
        "על התשובות להכיל את חישובי העזר והנימוקים בצורה מסודרת. לא יינתן ניקוד מלא לתשובות חלקיות, גם אם הן נכונות.",
        "בהצלחה!",
    ]
    for i, line in enumerate(lines):
        add_rtl_paragraph(doc, line, bold=i in {0, 3, 6, 11}, size=12 if i in {0, 3} else 10.5, space_after=2)
    add_rtl_paragraph(doc, space_after=8)


def add_datasets(doc: Document) -> None:
    add_heading(doc, "אוסף נתונים לבחינה", level=1)
    for dataset in DATASETS:
        add_rtl_paragraph(doc, dataset["title"], bold=True, size=11.5, space_after=2)
        add_rtl_paragraph(doc, dataset["description"], size=10.5, space_after=4)
        for table in dataset["tables"]:
            add_data_table(doc, table["headers"], table["rows"])


def add_question(doc: Document, question: dict) -> None:
    add_heading(doc, question["title"], level=2)
    add_rtl_paragraph(doc, question["intro"], size=10.5, space_after=4)
    num_id = create_list_numbering(doc, ordered=True, rtl=True)
    for part in question["parts"]:
        paragraph = add_rtl_paragraph(doc)
        set_paragraph_numbering(paragraph, num_id)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.right_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-18)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(part)
        run.font.name = "Arial"
        set_run_size(run, 10.8)
        set_run_direction(run, True)
    question_number = question["title"].split(" ")[1]
    add_rtl_paragraph(doc, f"פתרון שאלה {question_number} - יש להציג את הפתרון כאן", bold=True, size=10.5, space_after=10)


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_exam_header(doc)
    add_datasets(doc)
    for question in QUESTIONS:
        add_question(doc, question)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
