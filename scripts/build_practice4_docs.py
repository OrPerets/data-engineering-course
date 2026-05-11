from __future__ import annotations

import re
import textwrap
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/orperetz/Documents/shenkar/הנדסת נתונים/קורס")
EXERCISES = ROOT / "exercises"
BUILD = ROOT / "build"
CONTENT_WIDTH_CM = 17.6
COMPACT_TABLE_WIDTH_CM = 15.2
NARROW_TABLE_WIDTH_CM = 12.8
DXA_PER_CM = 1440 / 2.54

TITLE_COLOR = RGBColor(42, 82, 125)
ACCENT_FILL = "D9EAF7"
NOTE_FILL = "F2F6FA"
CODE_FILL = "F4F6F8"
BORDER_COLOR = "A9BACB"
DIAGRAM_BG = (248, 250, 252)
DIAGRAM_TEXT = (32, 44, 63)
DIAGRAM_ACCENT = (42, 82, 125)
DIAGRAM_LIGHT = (217, 234, 247)
DIAGRAM_MID = (235, 242, 248)

HEBREW_RE = re.compile(r"[\u0590-\u05FF]")


def insert_before_any(parent, child, before_tags: tuple[str, ...]) -> None:
    before_qnames = {qn(tag) for tag in before_tags}
    for idx, existing in enumerate(list(parent)):
        if existing.tag in before_qnames:
            parent.insert(idx, child)
            return
    parent.append(child)


def has_hebrew(text: str) -> bool:
    return bool(HEBREW_RE.search(text or ""))


def set_paragraph_direction(paragraph, rtl: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        insert_before_any(
            p_pr,
            bidi,
            (
                "w:adjustRightInd",
                "w:snapToGrid",
                "w:spacing",
                "w:ind",
                "w:contextualSpacing",
                "w:mirrorIndents",
                "w:suppressOverlap",
                "w:jc",
                "w:textDirection",
                "w:textAlignment",
                "w:textboxTightWrap",
                "w:outlineLvl",
                "w:divId",
                "w:cnfStyle",
                "w:rPr",
                "w:sectPr",
            ),
        )
    bidi.set(qn("w:val"), "1" if rtl else "0")
    if rtl:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_section_direction(section, rtl: bool = True) -> None:
    if not rtl:
        return
    sect_pr = section._sectPr
    bidi = sect_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        insert_before_any(sect_pr, bidi, ("w:rtlGutter", "w:docGrid", "w:printerSettings"))
    bidi.set(qn("w:val"), "1")

    rtl_gutter = sect_pr.find(qn("w:rtlGutter"))
    if rtl_gutter is None:
        rtl_gutter = OxmlElement("w:rtlGutter")
        insert_before_any(sect_pr, rtl_gutter, ("w:docGrid", "w:printerSettings"))
    rtl_gutter.set(qn("w:val"), "1")


def set_run_direction(run, rtl: bool) -> None:
    if not rtl:
        return
    r_pr = run._r.get_or_add_rPr()
    rtl_elem = r_pr.find(qn("w:rtl"))
    if rtl_elem is None:
        rtl_elem = OxmlElement("w:rtl")
        rtl_elem.set(qn("w:val"), "1")
        r_pr.append(rtl_elem)


def set_run_size(run, size_pt: float) -> None:
    run.font.size = Pt(size_pt)
    r_pr = run._r.get_or_add_rPr()
    half_points = str(int(round(size_pt * 2)))
    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), half_points)


def set_style_direction(style, rtl: bool, align: str = "left") -> None:
    style_el = style.element
    p_pr = style_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        insert_before_any(style_el, p_pr, ("w:rPr", "w:tblPr", "w:trPr", "w:tcPr"))
    if rtl:
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            insert_before_any(
                p_pr,
                bidi,
                (
                    "w:adjustRightInd",
                    "w:snapToGrid",
                    "w:spacing",
                    "w:ind",
                    "w:contextualSpacing",
                    "w:mirrorIndents",
                    "w:suppressOverlap",
                    "w:jc",
                    "w:textDirection",
                    "w:textAlignment",
                    "w:textboxTightWrap",
                    "w:outlineLvl",
                    "w:divId",
                    "w:cnfStyle",
                    "w:rPr",
                    "w:sectPr",
                ),
            )
        bidi.set(qn("w:val"), "1")
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        insert_before_any(
            p_pr,
            jc,
            (
                "w:textDirection",
                "w:textAlignment",
                "w:textboxTightWrap",
                "w:outlineLvl",
                "w:divId",
                "w:cnfStyle",
                "w:pPrChange",
            ),
        )
    jc.set(qn("w:val"), align)

    r_pr = style_el.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        style_el.append(r_pr)
    if rtl:
        rtl_el = r_pr.find(qn("w:rtl"))
        if rtl_el is None:
            rtl_el = OxmlElement("w:rtl")
            r_pr.append(rtl_el)
        rtl_el.set(qn("w:val"), "1")


def configure_document_defaults(doc: Document) -> None:
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)

    p_pr_default = doc_defaults.find(qn("w:pPrDefault"))
    if p_pr_default is None:
        p_pr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(p_pr_default)
    p_pr = p_pr_default.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_pr_default.append(p_pr)
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        insert_before_any(
            p_pr,
            bidi,
            (
                "w:adjustRightInd",
                "w:snapToGrid",
                "w:spacing",
                "w:ind",
                "w:contextualSpacing",
                "w:mirrorIndents",
                "w:suppressOverlap",
                "w:jc",
                "w:textDirection",
                "w:textAlignment",
                "w:textboxTightWrap",
                "w:outlineLvl",
                "w:divId",
                "w:cnfStyle",
                "w:rPr",
                "w:sectPr",
            ),
        )
    bidi.set(qn("w:val"), "1")
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        insert_before_any(
            p_pr,
            jc,
            (
                "w:textDirection",
                "w:textAlignment",
                "w:textboxTightWrap",
                "w:outlineLvl",
                "w:divId",
                "w:cnfStyle",
                "w:pPrChange",
            ),
        )
    jc.set(qn("w:val"), "right")

    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        insert_before_any(
            tc_pr,
            shd,
            (
                "w:noWrap",
                "w:tcMar",
                "w:textDirection",
                "w:tcFitText",
                "w:vAlign",
                "w:hideMark",
                "w:headers",
                "w:cellIns",
                "w:cellDel",
                "w:cellMerge",
                "w:tcPrChange",
            ),
        )
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph, color=BORDER_COLOR, size=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        insert_before_any(
            p_pr,
            p_bdr,
            (
                "w:shd",
                "w:tabs",
                "w:suppressAutoHyphens",
                "w:kinsoku",
                "w:wordWrap",
                "w:overflowPunct",
                "w:topLinePunct",
                "w:autoSpaceDE",
                "w:autoSpaceDN",
                "w:bidi",
                "w:adjustRightInd",
                "w:snapToGrid",
                "w:spacing",
                "w:ind",
                "w:contextualSpacing",
                "w:mirrorIndents",
                "w:suppressOverlap",
                "w:jc",
                "w:textDirection",
                "w:textAlignment",
                "w:textboxTightWrap",
                "w:outlineLvl",
                "w:divId",
                "w:cnfStyle",
                "w:rPr",
                "w:sectPr",
            ),
        )
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def cm_to_dxa(value_cm: float) -> int:
    return int(round(value_cm * DXA_PER_CM))


def set_table_width(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        insert_before_any(tbl_pr, tbl_layout, ("w:tblLook", "w:tblCaption", "w:tblDescription", "w:tblPrChange"))
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_direction(table, rtl: bool = True) -> None:
    if not rtl:
        return
    tbl_pr = table._tbl.tblPr
    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        insert_before_any(
            tbl_pr,
            bidi_visual,
            (
                "w:tblStyleRowBandSize",
                "w:tblStyleColBandSize",
                "w:tblW",
                "w:jc",
                "w:tblCellSpacing",
                "w:tblInd",
                "w:tblBorders",
                "w:shd",
                "w:tblLayout",
                "w:tblCellMar",
                "w:tblLook",
                "w:tblCaption",
                "w:tblDescription",
                "w:tblPrChange",
            ),
        )
    bidi_visual.set(qn("w:val"), "1")
    table.table_direction = WD_TABLE_DIRECTION.RTL


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def next_numbering_id(numbering, tag: str, attr: str) -> int:
    values = []
    for element in numbering.findall(qn(tag)):
        raw_value = element.get(qn(attr))
        if raw_value is not None:
            values.append(int(raw_value))
    return (max(values) + 1) if values else 1


def create_list_numbering(doc: Document, ordered: bool, rtl: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_num_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level_type)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "right" if rtl else "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    if rtl:
        ind.set(qn("w:right"), "720")
    else:
        ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)

    if not ordered:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        lvl.append(r_pr)

    abstract_num.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract_num)
    else:
        numbering.insert(list(numbering).index(first_num), abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_paragraph_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_style = p_pr.find(qn("w:pStyle"))
        insert_at = list(p_pr).index(p_style) + 1 if p_style is not None else 0
        p_pr.insert(insert_at, num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))

    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    set_section_direction(section, rtl=True)
    configure_document_defaults(doc)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    set_style_direction(normal, rtl=True, align="left")

    for style_name, font_size, bold in (
        ("Heading 1", 17, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.color.rgb = TITLE_COLOR
        set_style_direction(style, rtl=True, align="left")

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(9)
    set_style_direction(styles["CodeBlock"], rtl=False, align="left")

    if "Footer" in styles:
        set_style_direction(styles["Footer"], rtl=True, align="center")

    for list_style_name in ("List Paragraph", "List Bullet", "List Number"):
        try:
            list_style = styles[list_style_name]
        except KeyError:
            continue
        list_style.font.name = "Arial"
        list_style.font.size = Pt(11)
        set_style_direction(list_style, rtl=True, align="left")

    footer = section.footer
    fp = footer.paragraphs[0]
    set_paragraph_direction(fp, True)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)
    spacer = fp.add_run(" ")
    spacer.font.name = "Arial"
    spacer.font.size = Pt(9)
    set_run_direction(spacer, True)
    fp_run = fp.add_run("עמוד")
    fp_run.font.name = "Arial"
    fp_run.font.size = Pt(9)
    set_run_direction(fp_run, True)

    settings_el = doc.settings.element
    zoom = settings_el.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")
    theme_font_lang = settings_el.find(qn("w:themeFontLang"))
    if theme_font_lang is not None:
        theme_font_lang.set(qn("w:val"), "he-IL")
        theme_font_lang.set(qn("w:bidi"), "he-IL")


def apply_content_fixes(text: str, kind: str) -> str:
    if kind == "student":
        text = text.replace(
            "2. סווגו כל טבלת מקור כטבלה mutable, append-only או late-arriving, והסבירו בקצרה.",
            "2. סווגו כל טבלת מקור כטבלה mutable, append-only ו/או late-arriving, והסבירו בקצרה.",
        )
        text = text.replace(
            "2. שיעור מימוש הזמנות שבועי לפי ערוץ, המוגדר כ-`completed / total orders`.",
            "2. שיעור מימוש הזמנות שבועי לפי ערוץ, המוגדר כ-`completed / total orders` על בסיס הסטטוס האחרון לאחר dedup של `orders_src`.",
        )
    else:
        text = text.replace(
            "| `returns_src` | Late-arriving / append-only | אירועי החזרה מתווספים לאחר ההזמנה ולעיתים מגיעים בעיכוב. |",
            "| `returns_src` | Append-only עם אופי late-arriving | אירועי החזרה מתווספים לאחר ההזמנה ולעיתים מגיעים בעיכוב. |",
        )
        text = text.replace(
            "| fact_sales | customer_key | orders_src.customer_id | lookup ב-`dim_customer` לפי חלון SCD2 תקף | FK | reject |",
            "| fact_sales | customer_key | `orders_src.customer_id` + `orders_src.order_ts` | lookup ב-`dim_customer` לפי חלון SCD2 התקף לזמן ההזמנה | FK | reject |",
        )
        text = text.replace(
            "  LEFT JOIN (\n    SELECT order_id, line_id, SUM(return_amount) AS return_amount\n    FROM staging.returns_src\n    GROUP BY order_id, line_id\n  ) r\n    ON r.order_id = oi.order_id\n   AND r.line_id = oi.line_id\n) s",
            "  LEFT JOIN (\n    SELECT order_id, line_id, SUM(return_amount) AS return_amount\n    FROM staging.returns_src\n    GROUP BY order_id, line_id\n  ) r\n    ON r.order_id = oi.order_id\n   AND r.line_id = oi.line_id\n  WHERE o.order_status IN ('completed', 'refunded')\n) s",
        )
        text = text.replace(
            "### ג2.7 - טיפול בנתונים מאוחרים מ-`returns_src`\n",
            "הנחת עבודה: הזמנות במצב `cancelled` אינן נטענות ל-`fact_sales`, בעוד שהזמנות `refunded` נשמרות ב-fact וההחזרה מתקנת את `net_revenue`.\n\n### ג2.7 - טיפול בנתונים מאוחרים מ-`returns_src`\n",
        )
        text = text.replace(
            "GROUP BY p.product_name",
            "GROUP BY p.product_key, p.product_name",
        )
    return text


def add_cover(doc: Document, title: str, subtitle: str, note: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_direction(p, True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = TITLE_COLOR
    r.font.name = "Arial"
    set_run_direction(r, True)

    p2 = doc.add_paragraph()
    set_paragraph_direction(p2, True)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = TITLE_COLOR
    r2.font.name = "Arial"
    set_run_direction(r2, True)

    note_table = doc.add_table(rows=1, cols=1)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    note_table.allow_autofit = True
    set_table_direction(note_table, rtl=True)
    cell = note_table.cell(0, 0)
    shade_cell(cell, NOTE_FILL)
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    p3 = cell.paragraphs[0]
    set_paragraph_direction(p3, True)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = p3.add_run(note)
    rr.font.size = Pt(10.5)
    rr.font.name = "Arial"
    set_run_direction(rr, True)

    doc.add_paragraph()


def get_diagram_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, subtitle: str, fill: tuple[int, int, int]) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=DIAGRAM_ACCENT, width=3)
    title_font = get_diagram_font(30, bold=True)
    subtitle_font = get_diagram_font(21, bold=False)
    x0, y0, x1, y1 = xy
    title_bbox = draw.multiline_textbbox((0, 0), title, font=title_font, spacing=6, align="center")
    subtitle_bbox = draw.multiline_textbbox((0, 0), subtitle, font=subtitle_font, spacing=4, align="center")
    title_h = title_bbox[3] - title_bbox[1]
    subtitle_h = subtitle_bbox[3] - subtitle_bbox[1]
    total_h = title_h + subtitle_h + 18
    text_y = y0 + ((y1 - y0 - total_h) // 2)
    draw.multiline_text(((x0 + x1) / 2, text_y), title, font=title_font, fill=DIAGRAM_TEXT, anchor="ma", align="center", spacing=6)
    draw.multiline_text(((x0 + x1) / 2, text_y + title_h + 18), subtitle, font=subtitle_font, fill=DIAGRAM_TEXT, anchor="ma", align="center", spacing=4)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill=DIAGRAM_ACCENT, width=6)
    arrow_size = 12
    ex, ey = end
    draw.polygon(
        [(ex, ey), (ex - arrow_size * 2, ey - arrow_size), (ex - arrow_size * 2, ey + arrow_size)],
        fill=DIAGRAM_ACCENT,
    )


def draw_table_card(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    rows: list[str],
    fill: tuple[int, int, int],
) -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=16, fill=(255, 255, 255), outline=DIAGRAM_ACCENT, width=3)
    draw.rounded_rectangle((x0, y0, x1, y0 + 46), radius=16, fill=fill, outline=DIAGRAM_ACCENT, width=3)
    draw.rectangle((x0, y0 + 28, x1, y0 + 46), fill=fill, outline=fill)

    title_font = get_diagram_font(22, bold=True)
    row_font = get_diagram_font(18)
    draw.text(((x0 + x1) / 2, y0 + 23), title, font=title_font, fill=DIAGRAM_TEXT, anchor="ma")

    usable_top = y0 + 58
    row_height = max(24, (y1 - usable_top - 16) // max(len(rows), 1))
    for idx, row in enumerate(rows):
        y = usable_top + idx * row_height
        if idx > 0:
            draw.line([(x0 + 10, y - 6), (x1 - 10, y - 6)], fill=(221, 229, 236), width=2)
        draw.text((x0 + 16, y), row, font=row_font, fill=DIAGRAM_TEXT, anchor="la")


def draw_problem_panel(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, bullets: list[str]) -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=18, fill=DIAGRAM_MID, outline=DIAGRAM_ACCENT, width=3)
    title_font = get_diagram_font(28, bold=True)
    row_font = get_diagram_font(21)
    draw.text(((x0 + x1) / 2, y0 + 28), title, font=title_font, fill=DIAGRAM_TEXT, anchor="ma")
    start_y = y0 + 76
    row_gap = max(30, (y1 - start_y - 20) // max(len(bullets), 1))
    for idx, bullet in enumerate(bullets):
        y = start_y + idx * row_gap
        draw.ellipse((x0 + 18, y - 4, x0 + 30, y + 8), fill=DIAGRAM_ACCENT)
        draw.text((x0 + 42, y), bullet, font=row_font, fill=DIAGRAM_TEXT, anchor="la")


def build_overview_diagram(output_path: Path) -> Path:
    image = Image.new("RGB", (1800, 1060), DIAGRAM_BG)
    draw = ImageDraw.Draw(image)

    header_font = get_diagram_font(38, bold=True)
    sub_font = get_diagram_font(22)
    draw.text((900, 48), "Source Tables and Problem Framing", font=header_font, fill=DIAGRAM_ACCENT, anchor="ma")
    draw.text(
        (900, 96),
        "UML-like view of the source tables followed by the core data-quality problem",
        font=sub_font,
        fill=DIAGRAM_TEXT,
        anchor="ma",
    )

    source_tables = [
        ((70, 160, 390, 420), "orders_src", ["order_id", "customer_ref", "order_ts", "channel", "order_status", "updated_at"], DIAGRAM_LIGHT),
        ((420, 160, 740, 420), "order_items_src", ["order_id", "line_id", "product_id", "quantity", "unit_price", "discount_amount", "updated_at"], DIAGRAM_LIGHT),
        ((770, 160, 1090, 420), "customers_src", ["customer_ref", "full_name", "segment", "city", "region", "updated_at"], DIAGRAM_LIGHT),
        ((1120, 160, 1440, 420), "products_src", ["product_id", "product_name", "category", "brand", "supplier_id", "updated_at"], DIAGRAM_LIGHT),
        ((525, 460, 875, 700), "returns_src", ["return_id", "order_id", "line_id", "return_ts", "return_amount", "reason_code", "ingested_at"], DIAGRAM_LIGHT),
    ]
    for xy, title, rows, fill in source_tables:
        draw_table_card(draw, xy, title, rows, fill)

    connector_y = 744
    draw.line([(232, 420), (232, connector_y), (900, connector_y)], fill=DIAGRAM_ACCENT, width=5)
    draw.line([(582, 420), (582, connector_y)], fill=DIAGRAM_ACCENT, width=5)
    draw.line([(932, 420), (932, connector_y)], fill=DIAGRAM_ACCENT, width=5)
    draw.line([(1282, 420), (1282, connector_y), (900, connector_y)], fill=DIAGRAM_ACCENT, width=5)
    draw.line([(700, 700), (700, connector_y)], fill=DIAGRAM_ACCENT, width=5)
    draw_arrow(draw, (900, connector_y), (900, 812))

    problem_box = (160, 804, 1640, 1000)
    draw_problem_panel(
        draw,
        problem_box,
        "Problem to Solve",
        [
            "The same customer can appear with different identifiers across systems",
            "channel / segment / region / category values are not standardized",
            "orders can be updated, deduplicated, or cancelled after the first load",
            "returns can arrive late and distort revenue KPIs if handled incorrectly",
        ],
    )

    legend_font = get_diagram_font(18)
    draw.rounded_rectangle((90, 114, 1710, 144), radius=10, fill=(241, 245, 249), outline=(196, 208, 220), width=1)
    draw.text(
        (900, 129),
        "This diagram intentionally stops at the source landscape and the problem statement.",
        font=legend_font,
        fill=DIAGRAM_TEXT,
        anchor="ma",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)
    return output_path


def add_overview_section(doc: Document, diagram_path: Path) -> None:
    p = doc.add_paragraph(style="Heading 1")
    set_paragraph_direction(p, True)
    run = p.add_run("מבט על התרגול")
    set_run_direction(run, True)
    set_paragraph_bottom_border(p, size=4)

    intro = doc.add_paragraph()
    set_paragraph_direction(intro, True)
    rr = intro.add_run(
        "התרשים הבא מציג את טבלאות המקור ואת בעיות הנתונים שהן יוצרות. מטרתו למקד את הסטודנטים בהבנת הבעיה העסקית והטכנית לפני תכנון הפתרון."
    )
    rr.font.name = "Arial"
    rr.font.size = Pt(11)
    set_run_direction(rr, True)

    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(diagram_path), width=Cm(17.0))

    caption = doc.add_paragraph()
    set_paragraph_direction(caption, True)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run("איור 1: תצוגת UML של טבלאות המקור ולאחריה ניסוח ברור של בעיית הנתונים.")
    cap_run.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.name = "Arial"
    set_run_direction(cap_run, True)

    doc.add_paragraph()


def iter_blocks(soup: BeautifulSoup):
    for child in soup.contents:
        if isinstance(child, NavigableString) and not child.strip():
            continue
        if isinstance(child, Tag):
            yield child


def add_inline_runs(paragraph, node, rtl: bool | None = None) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if not text:
                continue
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            actual_rtl = has_hebrew(text) if rtl is None else rtl
            set_run_direction(run, actual_rtl)
        elif child.name == "strong":
            run = paragraph.add_run(child.get_text())
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
            actual_rtl = has_hebrew(child.get_text()) if rtl is None else rtl
            set_run_direction(run, actual_rtl)
        elif child.name == "em":
            run = paragraph.add_run(child.get_text())
            run.italic = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
            actual_rtl = has_hebrew(child.get_text()) if rtl is None else rtl
            set_run_direction(run, actual_rtl)
        elif child.name == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            set_run_direction(run, False)
        elif child.name == "br":
            paragraph.add_run("\n")
        else:
            add_inline_runs(paragraph, child, rtl=rtl)


def wrap_code(text: str, width: int = 78) -> str:
    wrapped_lines = []
    for line in text.splitlines():
        if not line.strip():
            wrapped_lines.append("")
            continue
        indent = len(line) - len(line.lstrip(" "))
        subsequent = " " * min(indent + 2, 12)
        chunks = textwrap.wrap(
            line,
            width=width,
            replace_whitespace=False,
            drop_whitespace=False,
            break_long_words=False,
            subsequent_indent=subsequent,
        )
        wrapped_lines.extend(chunks or [line])
    return "\n".join(wrapped_lines)


def add_code_block(doc: Document, code_text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_FILL)
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    p = cell.paragraphs[0]
    p.style = doc.styles["CodeBlock"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(wrap_code(code_text.rstrip()))
    run.font.name = "Courier New"
    run.font.size = Pt(8.8)
    doc.add_paragraph()


def estimate_column_widths(table_rows: list[list[str]], total_dxa: int) -> list[int]:
    col_count = max(len(row) for row in table_rows)
    lengths = [1] * col_count
    for row in table_rows:
        for i, value in enumerate(row):
            lengths[i] = max(lengths[i], min(max(len(value), 4), 28))
    total = sum(lengths)
    widths = [int(total_dxa * length / total) for length in lengths]
    widths[-1] += total_dxa - sum(widths)
    return widths


def table_profile(rows: list[list[str]]) -> tuple[int, float, int]:
    col_count = max(len(row) for row in rows)
    if col_count <= 2:
        return cm_to_dxa(11.4), 8.2, 25
    if col_count <= 4:
        return cm_to_dxa(13.2), 7.8, 24
    return cm_to_dxa(COMPACT_TABLE_WIDTH_CM), 6.0, 10


def add_markdown_table(doc: Document, table_tag: Tag) -> None:
    rows = []
    for tr in table_tag.find_all("tr"):
        row = []
        for cell in tr.find_all(["th", "td"]):
            row.append(cell.get_text(" ", strip=True))
        if row:
            rows.append(row)
    if not rows:
        return

    rtl = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_direction(table, rtl=True)

    total_width_dxa, font_size, margin = table_profile(rows)
    set_table_width(table, total_width_dxa)
    widths = estimate_column_widths(rows, total_width_dxa)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell, top=margin, start=margin, bottom=margin, end=margin)
            if r_idx == 0:
                shade_cell(cell, ACCENT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            cell_rtl = has_hebrew(value)
            set_paragraph_direction(p, cell_rtl)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Arial"
            set_run_size(run, font_size)
            run.bold = r_idx == 0
            set_run_direction(run, cell_rtl)
    doc.add_paragraph()


def add_list(doc: Document, list_tag: Tag, ordered: bool) -> None:
    for idx, li in enumerate(list_tag.find_all("li", recursive=False), start=1):
        text = li.get_text(" ", strip=True)
        p = doc.add_paragraph()
        set_paragraph_direction(p, True)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.right_indent = Cm(0.4)
        prefix_text = f"{idx}. " if ordered else "- "
        prefix = p.add_run(prefix_text)
        prefix.font.name = "Arial"
        prefix.font.size = Pt(11)
        set_run_direction(prefix, True)
        add_inline_runs(p, li, rtl=True)


def render_html(doc: Document, html: str, skip_first_h1: bool = True) -> None:
    soup = BeautifulSoup(html, "lxml")
    first_h1_skipped = False

    for block in iter_blocks(soup.body or soup):
        name = block.name.lower()
        if name == "h1":
            if skip_first_h1 and not first_h1_skipped:
                first_h1_skipped = True
                continue
            p = doc.add_paragraph(style="Heading 1")
            set_paragraph_direction(p, True)
            add_inline_runs(p, block, rtl=True)
        elif name == "h2":
            p = doc.add_paragraph(style="Heading 1")
            set_paragraph_direction(p, True)
            add_inline_runs(p, block, rtl=True)
            set_paragraph_bottom_border(p, size=4)
        elif name == "h3":
            p = doc.add_paragraph(style="Heading 2")
            set_paragraph_direction(p, True)
            add_inline_runs(p, block, rtl=True)
        elif name == "h4":
            p = doc.add_paragraph(style="Heading 3")
            set_paragraph_direction(p, True)
            add_inline_runs(p, block, rtl=True)
        elif name == "p":
            text = block.get_text(" ", strip=False)
            if not text.strip():
                continue
            p = doc.add_paragraph()
            set_paragraph_direction(p, True)
            add_inline_runs(p, block, rtl=True)
        elif name == "ul":
            add_list(doc, block, ordered=False)
        elif name == "ol":
            add_list(doc, block, ordered=True)
        elif name == "pre":
            code = block.get_text()
            add_code_block(doc, code)
        elif name == "table":
            add_markdown_table(doc, block)
        elif name == "hr":
            doc.add_paragraph()


def build_one(
    source_name: str,
    output_name: str,
    kind: str,
    title: str,
    subtitle: str,
    note: str,
    mirror_output_name: str | None = None,
) -> None:
    source_path = EXERCISES / source_name
    output_path = EXERCISES / output_name

    text = source_path.read_text(encoding="utf-8")
    text = apply_content_fixes(text, kind)
    html = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
        output_format="html5",
    )

    doc = Document()
    configure_document(doc)
    add_cover(doc, title=title, subtitle=subtitle, note=note)
    if kind == "student":
        diagram_path = build_overview_diagram(EXERCISES / "practice4_overview.png")
        add_overview_section(doc, diagram_path)
    render_html(doc, html)
    doc.save(output_path)
    if mirror_output_name:
        mirror_path = BUILD / mirror_output_name
        mirror_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(mirror_path)


def main() -> None:
    build_one(
        source_name="practice4.md",
        output_name="practice4.docx",
        kind="student",
        title="תרגול 4: תכנון מחסן נתונים ותהליך טעינה",
        subtitle="מסמך תרגול מסודר לסטודנטים",
        note="המסמך מסודר לפי זרימת עבודה מומלצת: הבנת התרחיש, מקורות הנתונים, המידול, ה-ETL, ולבסוף השאילתות האנליטיות.",
        mirror_output_name="practice4.docx",
    )
    build_one(
        source_name="practice5.md",
        output_name="practice4solution.docx",
        kind="solution",
        title="תרגול 4: תכנון מחסן נתונים ותהליך טעינה - פתרון מוצע",
        subtitle="פתרון מסודר לפי אותו מבנה של דף התרגול",
        note="זהו פתרון לדוגמה. ייתכנו פתרונות נוספים שהם נכונים וסבירים, כל עוד ההנחות והלוגיקה מנומקות היטב.",
    )


if __name__ == "__main__":
    main()
